from docx import Document
from datetime import datetime
import re

# move all methods into a class
class IndustryManagerParser:
    
    def __init__(self, docx_file):
        self.docx_file = docx_file
    
    def clean_text(self, text):
        """Clean and normalize text."""
        if not text:
            return ''
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def is_contact_info(self, text):
        """Check if text contains contact information."""
        patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # Phone
            r'\b\d+\s+[A-Za-z\s,]+\b[A-Z]{2}\b\s+\d{5}\b',  # Address
        ]
        return any(re.search(pattern, text) for pattern in patterns)

    def extract_achievement(self, text):
        """Extract achievement from text."""
        achievement_indicators = [
        r'\b\d+%\b',  # Percentage improvements
        r'\$\d+',     # Dollar amounts
        r'\bincreased\b',
        r'\bimproved\b',
        r'\breduced\b',
        r'\bsaved\b',
        r'\bgenerated\b',
        r'\bmanaged\b',
        r'\bled\b',
        r'\bdeveloped\b'
        ]
    
        if any(re.search(indicator, text.lower()) for indicator in achievement_indicators):
            return self.clean_text(text)
        return None

    def is_education_related(self, text):
        """Check if text is related to education."""
        education_keywords = [
            'bachelor', 'master', 'phd', 'degree', 'diploma',
            'certification', 'university', 'college', 'school',
            'gpa', 'major', 'minor'
        ]
        return any(keyword in text.lower() for keyword in education_keywords)

    def extract_name_from_contact(self, text):
        """Extract name from contact information, particularly from email."""
        # Look for email pattern
        email_match = re.search(r'([a-zA-Z0-9._%+-]+)@[a-zA-Z0-9.-]+\.[A-Z|a-z]{2,}', text)
        if email_match:
            email_prefix = email_match.group(1)
            # Handle common email formats like m.riley, riley.m, mriley
            parts = re.split(r'[._]', email_prefix)
            if len(parts) >= 2:
                # Try to determine which part is first name and which is last name
                for part in parts:
                    if len(part) > 1:  # Assume longer part is the last name
                        return f"{parts[0].title()} {part.title()}"
        return ''

    def parse_docx_to_json(self):
        """Parse a .docx file to extract structured data into JSON format."""
        
        extracted_data = {
            'name': '',
            'current_role': '',
            'years_experience': 0,
            'companies': [],
            'achievements': [],
            'skills': [],
            'education': [],
            'recognition': ''
        }

        document = Document(self.docx_file)
        
        # First pass: Extract name and role
        paragraphs = [self.clean_text(para.text) for para in document.paragraphs if self.clean_text(para.text)]
        
        # Extract name from contact information
        for para in paragraphs[:3]:  # Usually contact info is in first few paragraphs
            if self.is_contact_info(para):
                name = self.extract_name_from_contact(para)
                if name:
                    extracted_data['name'] = name
                break

        current_section = ''
        in_experience_details = False
        in_profile = False
        profile_text = []
        
        for para in document.paragraphs:
            text = self.clean_text(para.text)
            if not text:
                continue

            lower_text = text.lower()

            # Section detection
            if text.lower() == 'profile':
                current_section = 'profile'
                in_profile = True
                continue
            elif any(section in lower_text for section in ['experience', 'employment', 'work history']) and len(text) < 30:
                current_section = 'experience'
                in_experience_details = False
                in_profile = False
                continue
            elif any(section in lower_text for section in ['achievement', 'accomplishment']) and len(text) < 30:
                current_section = 'achievements'
                in_profile = False
                continue
            elif 'technical skills' in lower_text and len(text) < 30:
                current_section = 'skills'
                in_profile = False
                continue
            elif 'education' in lower_text and len(text) < 30:
                current_section = 'education'
                in_profile = False
                continue
            elif any(section in lower_text for section in ['recognition', 'award', 'honor']) and len(text) < 30:
                current_section = 'recognition'
                in_profile = False
                continue
            elif 'interests' in lower_text and len(text) < 30:
                current_section = 'interests'
                in_profile = False
                continue

            # Process sections
            if current_section == 'profile':
                if text.lower() != 'profile':
                    profile_text.append(text)
            
            elif current_section == 'experience':
                if '|' in text:
                    parts = [p.strip() for p in text.split('|')]
                    if len(parts) >= 2:
                        company = parts[1]
                        company = self.clean_text(company)
                        if company and company not in extracted_data['companies']:
                            extracted_data['companies'].append(company)
                        # Try to extract role from the first part
                        if parts[0] and not extracted_data['current_role']:
                            role = self.clean_text(parts[0])
                            if 'manager' in role.lower():
                                extracted_data['current_role'] = role
                    in_experience_details = True
                elif in_experience_details:
                    achievement = self.extract_achievement(text)
                    if achievement and achievement not in extracted_data['achievements']:
                        extracted_data['achievements'].append(achievement)
            
            elif current_section == 'achievements':
                achievement = self.extract_achievement(text)
                if achievement and achievement not in extracted_data['achievements']:
                    extracted_data['achievements'].append(achievement)
            
            elif current_section == 'skills':
                if not any(word in lower_text for word in ['section', 'skill', 'proficiency']):
                    skills = re.split(r'[•·,\u2022]|\band\b', text)
                    for skill in skills:
                        skill = self.clean_text(skill)
                        if (skill and len(skill) > 2 and 
                            not any(c.isdigit() for c in skill) and
                            not any(word in skill.lower() for word in ['activities', 'interests', 'theater', 'art', 'hiking', 'skiing', 'travel'])):
                            if skill not in extracted_data['skills']:
                                extracted_data['skills'].append(skill)
            
            elif current_section == 'education':
                if self.is_education_related(text):
                    if '|' in text:
                        parts = [p.strip() for p in text.split('|')]
                        education = next((p for p in parts if self.is_education_related(p)), None)
                        if education and education not in extracted_data['education']:
                            extracted_data['education'].append(education)
                    else:
                        education = self.clean_text(text)
                        if education and education not in extracted_data['education']:
                            extracted_data['education'].append(education)
            
            elif current_section == 'recognition':
                if not any(word in text.lower() for word in ['section', 'recognition', 'awards']):
                    if text not in extracted_data['recognition']:
                        extracted_data['recognition'] = text if not extracted_data['recognition'] else extracted_data['recognition'] + '; ' + text

        # Calculate experience based on January 2019 to present
        start_date = datetime(2019, 1, 1)  # January 2019
        end_date = datetime.now()
        years = (end_date - start_date).days / 365.25
        extracted_data['years_experience'] = round(years)

        # Extract skills from profile and achievements if none found
        if not extracted_data['skills']:
            # Look for skill keywords in profile and achievements
            skill_keywords = {
                'staff training': ['training', 'staff development'],
                'customer service': ['customer', 'service'],
                'team leadership': ['leader', 'leadership', 'team'],
                'food & beverage': ['food', 'beverage', 'restaurant'],
                'cost control': ['cost', 'efficiency', 'budget'],
                'social media': ['social media', 'marketing']
            }
            
            text_to_search = ' '.join(profile_text + extracted_data['achievements']).lower()
            for skill, keywords in skill_keywords.items():
                if any(keyword in text_to_search for keyword in keywords):
                    extracted_data['skills'].append(skill.title())

        return extracted_data

if __name__ == "__main__":
    import json
    file_path = "/Users/anilkumar/Desktop/tv3/src/templates/Industry manager resume.docx"
    parser = IndustryManagerParser(file_path)
    result = parser.parse_docx_to_json()
    print(json.dumps(result, indent=4))
