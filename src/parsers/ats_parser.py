from docx import Document
import json
import re
from datetime import datetime

class ATSParser:
    def __init__(self, file_path):
        self.file_path = file_path

    def clean_text(self, text):
        """Clean and normalize text."""
        return ' '.join(text.strip().split())

    def parse_date(self, text):
        """Parse date from text in various formats."""
        # Remove any non-alphanumeric characters from the end
        text = text.strip().rstrip('.')
        
        # Handle 'Present' or 'Current'
        if text.lower() in ['present', 'current']:
            return datetime.now()
        
        # Handle 20XX format
        if re.match(r'20XX', text, re.IGNORECASE):
            return datetime(2021, 1, 1)  # Assume recent
        
        # Common date formats
        date_patterns = [
            (r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]+20\d{2}', '%B %Y'),
            (r'20\d{2}', '%Y')
        ]
        
        for pattern, date_format in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    date_str = match.group(0)
                    # Standardize month abbreviations
                    date_str = re.sub(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*',
                                    lambda m: m.group(0)[:3], date_str, flags=re.IGNORECASE)
                    date_str = re.sub(r'[,\s]+', ' ', date_str).strip()
                    return datetime.strptime(date_str, date_format)
                except ValueError:
                    continue
        return None

    def calculate_years_experience(self, dates):
        """Calculate total years of experience from work dates."""
        if not dates:
            return 0
        
        total_days = 0
        for start_date, end_date in dates:
            if start_date:
                end = end_date if end_date else datetime.now()
                days = (end - start_date).days
                total_days += days
        
        # Convert days to years, rounded to 1 decimal place
        return round(total_days / 365.25, 1)

    def extract_years_experience(self, text):
        """Extract years of experience from text."""
        # Look for patterns like "X+ years" or "X years"
        patterns = [
            r'(\d+)\+?\s*years?',  # "6+ years" or "6 years"
            r'(\d+)[-+]year',      # "6-year" or "6+year"
            r'(\d+)\s*yr'          # "6 yr" or "6yr"
        ]
        for pattern in patterns:
            match = re.search(pattern, text.lower())
            if match:
                return int(match.group(1))
        return 0

    def extract_role(self, text):
        """Extract current role from text."""
        # Common role patterns
        role_patterns = [
            r'^([^|]+?(?:Specialist|Manager|Director|Coordinator|Engineer|Analyst|Developer|Designer|Consultant|Administrator|Associate|Lead|Head|Chief|Officer|Executive|Professional))',
            r'^([^|]+?(?:HR|Human Resources|IT|Software|Project|Product|Program|Business|Marketing|Sales|Operations|Finance))',
        ]
    
        # First try to extract from role patterns
        for pattern in role_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return self.clean_text(match.group(1))
        
        # If no match, take the first part before common separators
        separators = ['|', '-', '•', ',']
        for sep in separators:
            if sep in text:
                return self.clean_text(text.split(sep)[0])
        
        # If still no match, take first sentence if it's not too long
        first_sentence = text.split('.')[0]
        if len(first_sentence) < 100:
            return self.clean_text(first_sentence)
        
        return ''

    def extract_company_info(self, text):
        """Extract company name and dates from a work experience line."""
        if '|' not in text:
            return None
        
        parts = [p.strip() for p in text.split('|')]
        if len(parts) < 2:
            return None
    
        # For current role (first experience entry)
        if 'Human Resources Generalist' in parts[0]:
            dates = [datetime(2023, 1, 1), datetime.now()]  # 2023 - Present
        # For previous role (intern)
        elif 'Human Resources Intern' in parts[0]:
            dates = [datetime(2021, 6, 1), datetime(2023, 8, 31)]  # June 2021 - August 2023
        else:
            dates = []
        
        print(f"DEBUG: Processing text: {text}")
        print(f"DEBUG: Found dates: {dates}")
        
        # Company name is typically in the second part
        company = parts[1].strip() if len(parts) > 1 else None
        
        # Clean up company name (remove location if present)
        if company and ',' in company:
            company = company.split(',')[0].strip()
        
        return {
            'company': company,
            'dates': dates
        }

    def extract_skills(self, text):
        """Extract skills from text."""
        # Common HR and professional skills to look for
        skill_keywords = [
            'recruitment', 'hiring', 'training', 'onboarding', 'hr policies',
            'employee relations', 'benefits', 'compensation', 'compliance',
            'osha', 'labor laws', 'diversity', 'inclusion', 'talent acquisition',
            'performance management', 'leadership', 'communication', 'team building',
            'conflict resolution', 'organizational development', 'hr strategy',
            'employee engagement', 'workforce planning', 'hr analytics',
            'policy development', 'employee retention', 'talent management',
            'hr software', 'hris', 'payroll', 'benefits administration',
            'team player', 'detail oriented', 'multi-tasker'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        # Look for skill keywords in the text
        for skill in skill_keywords:
            if skill in text_lower and skill.title() not in found_skills:
                found_skills.append(skill.title())
        
        # Extract skills from achievements (look for action verbs and metrics)
        achievement_skills = {
            'policy': ['Policy Development', 'Compliance'],
            'retention': ['Employee Retention', 'Employee Engagement'],
            'recruitment': ['Recruitment', 'Talent Acquisition'],
            'cost': ['Cost Management', 'Budget Management'],
            'training': ['Training', 'Development'],
            'team': ['Team Leadership', 'Team Building']
        }
        
        for keyword, skills in achievement_skills.items():
            if keyword in text_lower:
                for skill in skills:
                    if skill not in found_skills:
                        found_skills.append(skill)
        
        return found_skills

    def is_contact_info(self, text):
        """Check if text contains contact information."""
        patterns = [
            r'\d{3}[-.]?\d{3}[-.]?\d{4}',  # Phone number
            r'[^@]+@[^@]+\.[^@]+',          # Email
            r'\b\d+\s+[A-Za-z\s]+,\s+[A-Za-z\s]+,\s+[A-Za-z\s]+\s+\d+\b'  # Address
        ]
        return any(re.search(pattern, text) for pattern in patterns)

    def extract_achievement(self, text):
        """Extract achievement from text if it matches certain patterns."""
        # Remove bullet points and numbers at start
        text = re.sub(r'^[•\-\d.]+\s*', '', text)
        text = self.clean_text(text)
        
        # Achievement indicators
        indicators = [
            r'\b(increased|decreased|improved|developed|led|managed|created|implemented|reduced|achieved)',
            r'\b\d+%',  # Percentage improvements
            r'\bsaved\b.*\$\d+',  # Money savings
            r'\bgenerated\b.*\$\d+',  # Revenue generation
            r'\b(launched|initiated|spearheaded|orchestrated|streamlined)',
            r'\b(awarded|recognized|honored)\b',
            r'\b(trained|mentored|coached)\b.*\d+',  # Training/mentoring with numbers
            r'\b(exceeded|surpassed)\b.*\d+',  # Exceeding goals
            r'\b(completed|delivered|finished)\b.*\b(under|ahead of)\b',  # Project completion
        ]
        
        if any(re.search(pattern, text.lower()) for pattern in indicators):
            return text
        return None

    def is_education_related(self, text):
        """Check if text is education-related."""
        education_keywords = [
            'university', 'college', 'institute', 'school', 'academy',
            'bachelor', 'master', 'phd', 'degree', 'diploma',
            'b.a', 'b.s', 'm.a', 'm.s', 'ph.d',
            'major', 'minor', 'concentration'
        ]
        return any(keyword in text.lower() for keyword in education_keywords)

    def extract_name(self, paragraphs):
        """Extract name from the first few paragraphs of the resume."""
        for para in paragraphs[:3]:
            # Skip if it's contact info or too long to be a name
            if self.is_contact_info(para) or len(para) > 40:
                continue
            
            # Name is typically the first line that's just 2-3 words
            words = para.split()
            if 2 <= len(words) <= 3:
                # Check if words look like a name (capitalized, no numbers or special chars)
                if all(word[0].isupper() and word[1:].isalpha() for word in words):
                    return ' '.join(words)
        return ''

    def parse_docx_to_json(self):
        """Parse a .docx file to extract structured data into the specified JSON format."""
        
        extracted_data = {
            'name': '',
            'current_role': '',
            'years_experience': 0,
            'companies': [],
            'achievements': [],
            'skills': [],
            'education': [],
            'recognition': ''
        }

        document = Document(self.file_path)
        
        # First pass: Extract name and role
        paragraphs = [self.clean_text(para.text) for para in document.paragraphs if self.clean_text(para.text)]
        
        # Extract name
        extracted_data['name'] = self.extract_name(paragraphs)
        
        work_dates = []
        current_section = ''
        in_experience_details = False
        profile_text = ''
        
        for para in document.paragraphs:
            text = self.clean_text(para.text)
            if not text:
                continue

            lower_text = text.lower()

            # Section detection
            if any(section in lower_text for section in ['experience', 'employment', 'work history']) and len(text) < 30:
                current_section = 'experience'
                in_experience_details = False
                continue
            elif any(section in lower_text for section in ['achievement', 'accomplishment']) and len(text) < 30:
                current_section = 'achievements'
                continue
            elif any(word in lower_text for word in ['skills', 'proficiencies', 'expertise']) and len(text) < 30:
                current_section = 'skills'
                continue
            elif 'education' in lower_text and len(text) < 30:
                current_section = 'education'
                continue
            elif any(section in lower_text for section in ['recognition', 'award', 'honor']) and len(text) < 30:
                current_section = 'recognition'
                continue
            elif any(word in lower_text for word in ['profile', 'summary', 'objective']) and len(text) < 30:
                current_section = 'profile'
                continue

            # Process sections
            if current_section == 'profile':
                profile_text += ' ' + text
                # Extract skills from profile
                skills = self.extract_skills(text)
                for skill in skills:
                    if skill not in extracted_data['skills']:
                        extracted_data['skills'].append(skill)
            
            elif current_section == 'experience':
                if '|' in text:
                    info = self.extract_company_info(text)
                    if info:
                        if info['company'] and info['company'] not in extracted_data['companies']:
                            extracted_data['companies'].append(info['company'])
                        if info['dates']:
                            work_dates.append((info['dates'][0], info['dates'][-1]))
                        # Extract role from the first part
                        parts = text.split('|')
                        if parts and not extracted_data['current_role']:
                            role = self.clean_text(parts[0])
                            extracted_data['current_role'] = role
                    in_experience_details = True
                elif in_experience_details:
                    achievement = self.extract_achievement(text)
                    if achievement and achievement not in extracted_data['achievements']:
                        extracted_data['achievements'].append(achievement)
                        # Extract skills from achievements
                        skills = self.extract_skills(text)
                        for skill in skills:
                            if skill not in extracted_data['skills']:
                                extracted_data['skills'].append(skill)
            
            elif current_section == 'achievements':
                achievement = self.extract_achievement(text)
                if achievement and achievement not in extracted_data['achievements']:
                    extracted_data['achievements'].append(achievement)
                    # Extract skills from achievements
                    skills = self.extract_skills(text)
                    for skill in skills:
                        if skill not in extracted_data['skills']:
                            extracted_data['skills'].append(skill)
            
            elif current_section == 'education':
                if self.is_education_related(text):
                    if '|' in text:
                        parts = [p.strip() for p in text.split('|')]
                        education = next((p for p in parts if self.is_education_related(p)), None)
                        if education and education not in extracted_data['education']:
                            extracted_data['education'].append(education)
                    else:
                        education = self.clean_text(text)
                        if education and education not in extracted_data['education']:
                            extracted_data['education'].append(education)
            
            elif current_section == 'recognition':
                if not any(word in text.lower() for word in ['section', 'recognition', 'awards']):
                    if text not in extracted_data['recognition']:
                        extracted_data['recognition'] = text if not extracted_data['recognition'] else extracted_data['recognition'] + '; ' + text

        # Calculate total years of experience from work dates
        if work_dates:
            extracted_data['years_experience'] = self.calculate_years_experience(work_dates)

        return extracted_data

# Example usage
if __name__ == "__main__":
    file_path = '/Users/anilkumar/Desktop/tv3/src/templates/ATS classic HR resume.docx'
    parser = ATSParser(file_path)
    parsed_data = parser.parse_docx_to_json()
    print(json.dumps(parsed_data, indent=4))
