import json
from docx import Document
import sys
import os
import re
from .base_resume_parser import BaseResumeParser
from typing import Optional

class DocumentParser():
    def __init__(self):
        self.name = "docx"

    def extract_contact_info(self, text):
        """Extract contact information from text."""
        parts = text.split('|')
        contact = {}
        for part in parts:
            part = part.strip()
            if '@' in part:
                contact['email'] = part
            elif 'linkedin.com' in part:
                contact['linkedin'] = part
            elif any(char.isdigit() for char in part) and '(' in part:
                contact['phone'] = part
            else:
                contact['address'] = part
        return contact

    def parse_position(self,text) -> Optional[dict]:
        """Extract position details from text."""
        parts = text.split('|')
        if len(parts) >= 3:
            return {
                "role": parts[0].strip(),
                "company": parts[1].strip(),
                "duration": parts[2].strip(),
                "key_achievements": []
            }
        return None


    def extract_docx_text(self,file_path):
        """Extract text from a DOCX file and return as structured JSON."""
        try:
            doc = Document(file_path)
            
            # Get all text
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Initialize structured data
            structured_data = {
                "personal_information": {
                    "name": "",
                    "contact": {},
                },
                "positions": [],
                "education": [],
                "skills": [],
                "interests": []
            }
            
            # Extract name from tables
            for table in doc.tables:
                if table.rows[0].cells[0].text.strip():
                    structured_data["personal_information"]["name"] = table.rows[0].cells[0].text.strip()
                    break
            
            # Process paragraphs
            current_section = None
            position_index = -1
            
            for i, text in enumerate(paragraphs):
                # Contact information
                if '|' in text and '@' in text:
                    structured_data["personal_information"]["contact"] = self.extract_contact_info(text.strip())
                
                # Position detection
                elif '|' in text and any(title in text.lower() for title in ['manager', 'director', 'supervisor', 'lead']):
                    position = self.parse_position(text)
                    if position:
                        structured_data["positions"].append(position)
                        position_index = len(structured_data["positions"]) - 1
                        current_section = "position"
                
                # Education detection
                elif any(degree in text for degree in ['B.S.', 'B.A.', 'M.S.', 'M.A.', 'Ph.D.', 'A.A.', 'A.S.', 'B.Tech.', 'MBA']):
                    if current_section == "education":
                        structured_data["education"].append(text)
                    else:
                        structured_data["education"] = [text]
                    current_section = "education"
                
                # Achievement collection for current position
                elif current_section == "position" and position_index >= 0:
                    if i > 0 and not any(marker in text.lower() for marker in ['education:', 'skills', 'profile', 'experience', 'activities']):
                        if not any(text.startswith(prefix) for prefix in ['Restaurant Manager', 'Director', 'Supervisor', 'Education']):
                            structured_data["positions"][position_index]["key_achievements"].append(text)
                
                # Skills
                elif "Skills & Abilities" in text:
                    current_section = "skills"
                    # Get skills from the next table
                    for table in doc.tables:
                        if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
                            skills = []
                            for cell in table.rows[0].cells:
                                skills.extend(cell.text.split('\n'))
                            structured_data["skills"] = [skill.strip() for skill in skills if skill.strip()]
                
                # Interests
                elif text.startswith("Theater") or "Activities and Interests" in text:
                    structured_data["interests"] = [interest.strip() for interest in text.split(',')]
            
            return structured_data
            
        except Exception as e:
            return {"error": str(e)}


if __name__ == "__main__":
    parser = DocumentParser()
    file_path = 'data/Industry manager resume.docx'

    if not os.path.exists(file_path):
        print(f"Error: File '{file_path}' does not exist.")
        sys.exit(1)

    # Extract the text
    result = parser.extract_docx_text(file_path)

    # Create output directory if it doesn't exist
    output_dir = 'output'
    os.makedirs(output_dir, exist_ok=True)

    # Save to JSON file
    output_file = os.path.join(output_dir, 'resume_data.json')
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        print(f"Data has been saved to {output_file}")
        
    except Exception as e:
        print(f"Error saving data: {e}")

    # Print the formatted JSON for verification
    print("\nExtracted Data:")

    print(json.dumps(result, indent=2, ensure_ascii=False))
